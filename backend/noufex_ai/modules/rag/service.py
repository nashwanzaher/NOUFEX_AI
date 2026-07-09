from __future__ import annotations

import json
import logging
import os
import tempfile
from pathlib import Path
from uuid import UUID

from sqlalchemy import func
from sqlalchemy.ext.asyncio import AsyncSession
from sqlmodel import select

from noufex_ai.exceptions import NotFoundError, UpstreamError, ValidationError
from noufex_ai.modules.agents.models import Document, KnowledgeChunk
from noufex_ai.modules.rag.schemas import (
    DocumentRead,
    RAGChunkResult,
    RAGSearchRequest,
    RAGSearchResponse,
)
from noufex_ai.settings import settings

logger = logging.getLogger(__name__)

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

_embedding_client = None


def _get_embedding_client():
    global _embedding_client
    if _embedding_client is None:
        from openai import AsyncOpenAI

        if not settings.openai_api_key:
            raise UpstreamError("OpenAI API key not configured")
        _embedding_client = AsyncOpenAI(
            api_key=settings.openai_api_key.get_secret_value(),
        )
    return _embedding_client


def _chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start = end - overlap
        if start + overlap >= len(text):
            break
    return chunks


async def _get_embeddings(texts: list[str]) -> list[list[float]]:
    from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=1, max=10),
        retry=retry_if_exception_type((Exception,)),
        reraise=True,
    )
    async def _get_embeddings_with_retry():
        client = _get_embedding_client()
        response = await client.embeddings.create(
            model=settings.openai_embedding_model,
            input=texts,
            dimensions=settings.openai_embedding_dimensions,
        )
        return [item.embedding for item in response.data]

    try:
        return await _get_embeddings_with_retry()
    except Exception as exc:
        logger.error("Embedding generation failed after retries: %s", exc)
        raise UpstreamError(f"Failed to generate embeddings after retries: {exc}") from exc


async def _extract_text(file_path: str, mime_type: str) -> str:
    if mime_type == "text/plain" or mime_type == "text/markdown":
        return Path(file_path).read_text(encoding="utf-8")
    elif mime_type == "text/html":
        from bs4 import BeautifulSoup

        html = Path(file_path).read_text(encoding="utf-8")
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n", strip=True)
    elif mime_type == "application/pdf":
        try:
            import pymupdf

            doc = pymupdf.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            raise ValidationError("PDF processing requires pymupdf. Install with: pip install pymupdf")
    elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        try:
            import docx

            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except ImportError:
            raise ValidationError("DOCX processing requires python-docx. Install with: pip install python-docx")
    else:
        raise ValidationError(f"Unsupported mime type: {mime_type}")


class RAGService:
    def __init__(self, session: AsyncSession) -> None:
        self.session = session

    async def upload_document(
        self,
        *,
        tenant_id: UUID,
        file_content: bytes,
        filename: str,
        mime_type: str,
    ) -> Document:
        if mime_type not in settings.allowed_upload_mime_types:
            raise ValidationError(f"Unsupported file type: {mime_type}")

        file_size = len(file_content)
        if file_size > settings.file_upload_max_size_mb * 1024 * 1024:
            raise ValidationError(f"File too large. Max size: {settings.file_upload_max_size_mb}MB")

        doc = Document(
            tenant_id=tenant_id,
            filename=f"{UUID.bytes.hex}-{filename}",
            original_filename=filename,
            mime_type=mime_type,
            file_size_bytes=file_size,
            status="processing",
        )
        self.session.add(doc)
        await self.session.flush()

        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                text = await _extract_text(tmp_path, mime_type)
                chunks = _chunk_text(text)

                if not chunks:
                    doc.status = "failed"
                    doc.error_message = "No content extracted from document"
                    self.session.add(doc)
                    await self.session.flush()
                    return doc

                embeddings = await _get_embeddings(chunks)

                for i, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
                    knowledge_chunk = KnowledgeChunk(
                        tenant_id=tenant_id,
                        document_id=doc.id,
                        chunk_index=i,
                        content=chunk_text,
                        embedding=embedding,
                        token_count=len(chunk_text.split()),
                    )
                    self.session.add(knowledge_chunk)

                doc.chunk_count = len(chunks)
                doc.status = "ready"
            finally:
                os.unlink(tmp_path)

        except Exception as exc:
            doc.status = "failed"
            doc.error_message = str(exc)[:1000]
            logger.error("Document processing failed for %s: %s", filename, exc)

        self.session.add(doc)
        await self.session.flush()
        return doc

    async def get_document(self, *, tenant_id: UUID, document_id: UUID) -> Document:
        result = await self.session.execute(
            select(Document).where(
                Document.id == document_id,
                Document.tenant_id == tenant_id,
                Document.deleted_at.is_(None),
            )
        )
        doc = result.scalar_one_or_none()
        if doc is None:
            raise NotFoundError(f"Document {document_id} not found")
        return doc

    async def list_documents(
        self, *, tenant_id: UUID, offset: int = 0, limit: int = 20
    ) -> tuple[list[Document], int]:
        base = select(Document).where(
            Document.tenant_id == tenant_id,
            Document.deleted_at.is_(None),
        )
        count_q = select(func.count()).select_from(base.subquery())
        count_result = await self.session.execute(count_q)
        total = count_result.scalar_one()

        q = base.order_by(Document.created_at.desc()).offset(offset).limit(limit)
        result = await self.session.execute(q)
        return list(result.scalars().all()), total

    async def delete_document(self, *, tenant_id: UUID, document_id: UUID) -> None:
        doc = await self.get_document(tenant_id=tenant_id, document_id=document_id)
        doc.soft_delete()
        self.session.add(doc)
        await self.session.flush()

    async def search(
        self, *, tenant_id: UUID, request: RAGSearchRequest
    ) -> RAGSearchResponse:
        query_embedding = await _get_embeddings([request.query])
        query_vec = query_embedding[0]

        # Use pgvector cosine distance operator (<=>) for efficient search
        # The <=> operator returns cosine distance (1 - cosine_similarity)
        cosine_dist = KnowledgeChunk.embedding.cosine_distance(query_vec)

        stmt = (
            select(
                KnowledgeChunk,
                cosine_dist.label("distance"),
            )
            .where(
                KnowledgeChunk.tenant_id == tenant_id,
                KnowledgeChunk.embedding.isnot(None),
            )
            .order_by(cosine_dist)
            .limit(request.top_k)
        )

        result = await self.session.execute(stmt)
        rows = result.all()

        results = []
        for chunk, distance in rows:
            score = 1.0 - distance  # convert distance to similarity
            if score >= request.score_threshold:
                meta = json.loads(chunk.metadata_json) if chunk.metadata_json else None
                results.append(
                    RAGChunkResult(
                        id=chunk.id,
                        document_id=chunk.document_id,
                        content=chunk.content,
                        score=round(score, 4),
                        metadata=meta,
                    )
                )

        return RAGSearchResponse(
            query=request.query,
            results=results,
            total=len(results),
        )


def _cosine_similarity(a: list[float], b: list[float]) -> float:
    """Compute cosine similarity between two vectors. Used in tests."""
    dot_product = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot_product / (norm_a * norm_b)
